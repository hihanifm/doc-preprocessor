"""
Run this script to generate sample_test_plan.docx in the current directory.
Usage: python create_sample.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_header_row(table, headers, bg="1F4E79"):
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = text
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell, bg)


def add_data_row(table, values):
    row = table.add_row()
    for i, text in enumerate(values):
        row.cells[i].text = text
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    return row


doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_heading("Sample Test Plan — User Management Module", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Version 1.0  |  Author: QA Team  |  Date: 2024-01-15"
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# ── Table of Contents (preamble — should be ignored) ──────────────────────────
doc.add_heading("Table of Contents", level=1)
doc.add_paragraph("1. Introduction")
doc.add_paragraph("2. Test Cases")
doc.add_paragraph("   2.1  TC_001_REGISTER")
doc.add_paragraph("   2.2  TC_002_LOGIN")
doc.add_paragraph("   2.3  TC_003_RESET_PWD")
doc.add_paragraph("   2.4  TC_004_UPDATE_PROFILE")
doc.add_paragraph("   2.5  TC_005_LOGOUT")
doc.add_paragraph("3. Appendix")
doc.add_page_break()

# ── Introduction (preamble — should be ignored) ────────────────────────────────
doc.add_heading("Introduction", level=1)
doc.add_paragraph(
    "This document covers end-to-end test cases for the User Management module. "
    "All test cases follow the naming convention <Title>_<ID>."
)
doc.add_paragraph(
    "Test environment: staging.example.com  |  Browser: Chrome 120  |  DB: PostgreSQL 15"
)
doc.add_page_break()

# ── Test Cases ─────────────────────────────────────────────────────────────────
doc.add_heading("Test Cases", level=1)


# ── TC 1: single table (Step | Expected Result) ────────────────────────────────
doc.add_heading("User Registration TC_001_REGISTER", level=2)
doc.add_paragraph(
    "Verify that a new user can successfully register with a valid email address "
    "and password. The system should create the account and send a confirmation email."
)

doc.add_paragraph("Preconditions:")
pre = doc.add_paragraph(style="List Bullet")
pre.add_run("The registration page is accessible at /register")
pre2 = doc.add_paragraph(style="List Bullet")
pre2.add_run("No existing account with the test email address exists")
pre3 = doc.add_paragraph(style="List Bullet")
pre3.add_run("SMTP service is configured and running")

doc.add_paragraph("")
doc.add_paragraph("Steps and Expected Results:").runs[0].font.bold = True

t1 = doc.add_table(rows=1, cols=3)
t1.style = "Table Grid"
add_header_row(t1, ["#", "Step", "Expected Result"])
steps_tc1 = [
    ("1", "Navigate to /register in a browser", "Registration form is displayed with fields: First Name, Last Name, Email, Password, Confirm Password"),
    ("2", "Enter valid first name 'John' and last name 'Doe'", "Fields are populated without validation errors"),
    ("3", "Enter a unique email address: newuser@example.com", "Email field accepts input; no error shown"),
    ("4", "Enter password 'P@ssw0rd123!' in both Password and Confirm Password fields", "Password strength indicator shows 'Strong'; fields match"),
    ("5", "Click the 'Register' button", "User is redirected to /dashboard; welcome banner shown; confirmation email sent to newuser@example.com"),
]
for row in steps_tc1:
    add_data_row(t1, list(row))

doc.add_paragraph("")


# ── TC 2: two separate tables (Steps table + Expected Results table) ───────────
doc.add_heading("User Login TC_002_LOGIN", level=2)
doc.add_paragraph(
    "Verify that a registered user can log in with correct credentials and is "
    "redirected to the dashboard. Invalid credentials should be rejected with an error."
)

doc.add_paragraph("Preconditions:")
for item in [
    "A registered account exists: testuser@example.com / P@ssw0rd123!",
    "The login page is accessible at /login",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_paragraph("")
doc.add_paragraph("Test Steps:").runs[0].font.bold = True

t2a = doc.add_table(rows=1, cols=2)
t2a.style = "Table Grid"
add_header_row(t2a, ["Step #", "Action"])
steps_tc2 = [
    ("1", "Open /login in a browser"),
    ("2", "Enter email: testuser@example.com"),
    ("3", "Enter password: P@ssw0rd123!"),
    ("4", "Click 'Login' button"),
    ("5", "Attempt login with wrong password: WrongPass!"),
]
for row in steps_tc2:
    add_data_row(t2a, list(row))

doc.add_paragraph("")
doc.add_paragraph("Expected Results:").runs[0].font.bold = True

t2b = doc.add_table(rows=1, cols=2)
t2b.style = "Table Grid"
add_header_row(t2b, ["Step #", "Expected Result"])
results_tc2 = [
    ("1", "Login form is displayed with Email and Password fields"),
    ("2", "Email field is populated"),
    ("3", "Password field shows masked input"),
    ("4", "User is redirected to /dashboard; username appears in nav bar"),
    ("5", "Error message displayed: 'Invalid email or password'; account not locked"),
]
for row in results_tc2:
    add_data_row(t2b, list(row))

doc.add_paragraph("")


# ── TC 3: single table ─────────────────────────────────────────────────────────
doc.add_heading("Password Reset TC_003_RESET_PWD", level=2)
doc.add_paragraph(
    "Verify that a user can reset their password via the forgot-password flow. "
    "A reset link should be emailed and the new password should work on next login."
)

doc.add_paragraph("Preconditions:")
for item in [
    "Registered account exists: testuser@example.com",
    "SMTP service is operational",
    "Password reset token expires after 60 minutes",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_paragraph("Applicability: All user roles (standard, admin)")
doc.add_paragraph("")
doc.add_paragraph("Steps and Expected Results:").runs[0].font.bold = True

t3 = doc.add_table(rows=1, cols=3)
t3.style = "Table Grid"
add_header_row(t3, ["#", "Test Step", "Expected Outcome"])
for row in [
    ("1", "Click 'Forgot Password' on /login", "Redirected to /forgot-password page"),
    ("2", "Enter registered email testuser@example.com and submit", "Success message shown; reset email sent within 30 seconds"),
    ("3", "Open the reset link from the email", "Password reset form displayed; link is valid"),
    ("4", "Enter new password 'NewP@ss456!' in both fields and submit", "Success confirmation shown; redirected to /login"),
    ("5", "Log in using testuser@example.com / NewP@ss456!", "Login succeeds; user reaches /dashboard"),
    ("6", "Attempt to reuse the old reset link", "Error: 'Reset link expired or already used'"),
]:
    add_data_row(t3, list(row))

doc.add_paragraph("")


# ── TC 4: two separate tables ──────────────────────────────────────────────────
doc.add_heading("Update User Profile TC_004_UPDATE_PROFILE", level=2)
doc.add_paragraph(
    "Verify that a logged-in user can update their profile information including "
    "display name, phone number, and avatar image. Changes should persist after page refresh."
)

doc.add_paragraph("Preconditions:")
for item in [
    "User is logged in as testuser@example.com",
    "Profile page is accessible at /profile",
    "Avatar upload accepts PNG and JPG under 2 MB",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_paragraph("")
doc.add_paragraph("Steps:").runs[0].font.bold = True

t4a = doc.add_table(rows=1, cols=2)
t4a.style = "Table Grid"
add_header_row(t4a, ["Step", "Action"])
for row in [
    ("1", "Navigate to /profile"),
    ("2", "Click 'Edit Profile'"),
    ("3", "Change Display Name to 'Johnny Doe'"),
    ("4", "Enter phone number: +1-555-000-1234"),
    ("5", "Upload a valid 800×800 PNG avatar"),
    ("6", "Click 'Save Changes'"),
    ("7", "Refresh the page"),
]:
    add_data_row(t4a, list(row))

doc.add_paragraph("")
doc.add_paragraph("Expected Results:").runs[0].font.bold = True

t4b = doc.add_table(rows=1, cols=2)
t4b.style = "Table Grid"
add_header_row(t4b, ["Step", "Expected Result"])
for row in [
    ("1", "Profile page displayed with current user data"),
    ("2", "Edit form becomes active; fields are editable"),
    ("3", "Display Name field shows 'Johnny Doe'"),
    ("4", "Phone field accepts E.164 formatted number"),
    ("5", "Avatar preview updates to the uploaded image"),
    ("6", "Success toast: 'Profile updated'; nav bar reflects new display name"),
    ("7", "All changes are persisted; avatar, name and phone are shown correctly"),
]:
    add_data_row(t4b, list(row))

doc.add_paragraph("")


# ── TC 5: single table ─────────────────────────────────────────────────────────
doc.add_heading("User Logout TC_005_LOGOUT", level=2)
doc.add_paragraph(
    "Verify that a logged-in user can log out successfully. The session should be "
    "invalidated and the user should be redirected to the login page. "
    "Accessing protected pages after logout should redirect back to /login."
)

doc.add_paragraph("Preconditions:")
for item in [
    "User is logged in and currently viewing /dashboard",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item)

doc.add_paragraph("")
doc.add_paragraph("Steps and Expected Results:").runs[0].font.bold = True

t5 = doc.add_table(rows=1, cols=3)
t5.style = "Table Grid"
add_header_row(t5, ["#", "Step", "Expected Result"])
for row in [
    ("1", "Click the user avatar / menu in the top-right corner", "Dropdown menu appears with options including 'Logout'"),
    ("2", "Click 'Logout'", "Session is cleared; user is redirected to /login with message 'You have been logged out'"),
    ("3", "Press the browser Back button", "Browser remains on /login or redirects back to it; /dashboard is not accessible"),
    ("4", "Manually navigate to /dashboard in the address bar", "Redirected to /login; dashboard content is not visible"),
    ("5", "Log in again with valid credentials", "Login succeeds normally; new session is created"),
]:
    add_data_row(t5, list(row))

doc.add_paragraph("")
doc.add_page_break()


# ── Appendix (suffix — should be ignored) ─────────────────────────────────────
doc.add_heading("Appendix", level=1)
doc.add_paragraph("A. Test Data Seed Script")
doc.add_paragraph("B. Environment Configuration")
doc.add_paragraph("C. Defect Log Template")

# ── Save ───────────────────────────────────────────────────────────────────────
out = "sample_test_plan.docx"
doc.save(out)
print(f"Created {out}")
